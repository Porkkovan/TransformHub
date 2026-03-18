"""
Generate four Word (.docx) documents for the TransformHub P0–P4 release:
  1. TransformHub-POV.docx
  2. TransformHub-CaseStudy-NationalBank.docx
  3. TransformHub-DemoGuide-P4.docx
  4. TransformHub-P0-P4-Enhancement-Summary.docx

Run from the docs/ directory:
    python generate_p4_documents.py
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = Path(__file__).parent


# ─── Shared rendering helpers (from convert_to_docx.py) ──────────────────────

def _shd(cell_or_para, fill_hex: str):
    """Apply fill colour to a table cell or paragraph."""
    if hasattr(cell_or_para, '_tc'):
        pr = cell_or_para._tc.get_or_add_tcPr()
    else:
        pr = cell_or_para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pr.append(shd)


def inline_format(para, text: str):
    """Add text with **bold** and `code` inline formatting."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        else:
            run = para.add_run(part)
            run.font.size = Pt(10)


def add_table_from_md(doc, lines, start_idx):
    """Parse a markdown table block and render as a Word table."""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1
    if len(table_lines) < 2:
        return start_idx

    headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
    data_rows = []
    for row_line in table_lines[2:]:
        cols = [c.strip() for c in row_line.split('|') if c.strip() != '']
        if cols:
            data_rows.append(cols)

    if not headers:
        return i

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = 'Table Grid'

    hdr_row = table.rows[0]
    for j, h in enumerate(headers[:num_cols]):
        cell = hdr_row.cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shd(cell, '1E3A8A')

    for ri, row_data in enumerate(data_rows):
        row = table.rows[ri + 1]
        for j in range(num_cols):
            cell = row.cells[j]
            text = row_data[j] if j < len(row_data) else ''
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            cell.text = text
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 1:
                _shd(cell, 'EFF6FF')

    doc.add_paragraph()
    return i


def convert_md_to_docx(md_path: Path, docx_path: Path):
    content = md_path.read_text(encoding='utf-8')
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_para.paragraph_format.left_indent = Inches(0.3)
                    code_para.paragraph_format.space_before = Pt(4)
                    code_para.paragraph_format.space_after = Pt(4)
                    for cl in code_lines:
                        run = code_para.add_run(cl + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                    _shd(code_para, 'EFF6FF')
                in_code_block = False
                code_lines = []
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() in ('---', '***', '___'):
            p = doc.add_paragraph('─' * 60)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            i = add_table_from_md(doc, lines, i)
            continue

        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x0C, 0x1C, 0x3A)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if line.startswith('## '):
            p = doc.add_heading(level=2)
            p.clear()
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        if line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:].strip())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            i += 1
            continue

        if re.match(r'^  [-*+] |^    [-*+] ', line):
            p = doc.add_paragraph(style='List Bullet 2')
            p.paragraph_format.left_indent = Inches(0.6)
            text = re.sub(r'^[\s]+[-*+] ', '', line).strip()
            inline_format(p, text)
            i += 1
            continue

        if re.match(r'^[-*+] ', line):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
            text = line[2:].strip()
            inline_format(p, text)
            i += 1
            continue

        m_num = re.match(r'^(\d+)\. ', line)
        if m_num:
            num = m_num.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            p.paragraph_format.space_after = Pt(3)
            # Write number as plain text so sequence is always preserved
            num_run = p.add_run(f"{num}. ")
            num_run.font.size = Pt(10)
            text = line[m_num.end():].strip()
            inline_format(p, text)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        inline_format(p, line.strip())
        i += 1

    doc.save(str(docx_path))
    print(f"  ✓  {docx_path.name}")


# ─── P0–P4 Enhancement Summary document ──────────────────────────────────────

ENHANCEMENTS = [
    # (phase, category, feature, description, file_or_component)
    ("P0", "Security", "Circuit Breaker",
     "CLOSED/OPEN/HALF_OPEN state machine on all LLM providers; automatic recovery with configurable failure thresholds and reset timeouts.",
     "app/services/circuit_breaker.py"),
    ("P0", "Identity", "SSO — Google OAuth",
     "Google OAuth 2.0 with domain-based organisation provisioning. Users from a registered domain are automatically mapped to the correct tenant.",
     "nextjs-app/src/lib/auth.ts"),
    ("P0", "Identity", "SSO — Azure Entra ID",
     "Azure Entra ID (OIDC) integration for enterprise customers. Supports domain provisioning and user sync.",
     "nextjs-app/src/lib/auth.ts"),
    ("P0", "Identity", "RBAC Hierarchy",
     "Four-level role hierarchy: VIEWER → ANALYST → ADMIN → SUPER_ADMIN. Permission inheritance enforced at API and UI layers.",
     "app/api/auth.py, nextjs-app middleware"),
    ("P0", "Identity", "API Key Management",
     "SHA-256 hashed API keys with org scope, configurable expiry, and scope arrays for programmatic access.",
     "app/api/api_keys.py"),
    ("P0", "Cost Control", "Per-Org LLM Budget Enforcement",
     "Monthly token and spend caps per organisation. Hard caps enforced pre-execution with 429 response. Budget period auto-resets via cron.",
     "app/services/budget_enforcement.py"),
    ("P0", "Security", "Prompt Injection Prevention",
     "10+ regex patterns detect jailbreak attempts in user-supplied text (agent input, document text, URL content) before LLM submission.",
     "nextjs-app/src/lib/api-validation.ts"),
    ("P0", "Reliability", "Rate Limiting",
     "Per-IP and per-org rate limiting on all API endpoints using sliding window counters in Redis.",
     "app/middleware/rate_limit.py"),
    ("P1", "Data Quality", "Tier 2a — Code Signal Extraction",
     "Regex-based extraction of timeout constants, SLA annotations, cron expressions, and ISO 8601 durations from uploaded code/architecture documents. Produces timing evidence with confidence 0.60–0.80.",
     "app/services/code_signal_extractor.py"),
    ("P1", "Data Quality", "Tier 3a — Jira Cycle Time Extraction",
     "Parses Jira issue changelogs to compute per-status dwell times. Maps status names to process_time / wait_time via configurable taxonomy. Produces jira_measured timings with confidence 0.70–0.95.",
     "app/services/jira_extractor.py"),
    ("P1", "Auditability", "Manual Timing Override",
     "Analysts can override any VSM step timing via UI. Writes timing_source='manual_override', confidence=1.0, and records previous value + override note in timing_overrides audit table.",
     "app/api/agents.py → /agents/timing-override"),
    ("P1", "Data Quality", "Agent Output Schema Validation",
     "Post-execution Pydantic validation of all 18 agent outputs. Non-critical validation failures emit warnings rather than failing the execution.",
     "app/agents/orchestrator.py"),
    ("P1", "Transparency", "Timing Provenance Badges",
     "Every VSM step carries timing_source (jira_measured / code_signals / manual_override / llm_estimated) and timing_confidence (0.0–1.0) surfaced in the UI.",
     "nextjs-app/src/components/vsm/"),
    ("P2", "DevOps", "GitHub Actions CI/CD",
     "Full pipeline: Ruff lint + pytest (PostgreSQL + Redis services), TypeScript type-check + ESLint + Next.js build, Jest, Docker multi-stage builds pushed to GHCR, npm audit + pip-audit.",
     ".github/workflows/"),
    ("P2", "Security", "Row-Level Security",
     "PostgreSQL RLS via app.current_org_id session variable. Data is physically inaccessible to other tenants even if application-layer bugs occur.",
     "migrations/rls.sql"),
    ("P2", "Performance", "HNSW Vector Index",
     "pgvector HNSW index on context_documents.embedding for sub-millisecond ANN search at scale. Replaces sequential scan on large knowledge bases.",
     "migrations/vector_index.sql"),
    ("P2", "Benchmarking", "Cross-Org Benchmark Aggregation",
     "Anonymised p25/p50/p75 flow efficiency benchmarks aggregated by industry type across all tenant organisations. Used to calibrate future state projections.",
     "app/services/benchmark_service.py"),
    ("P2", "Reliability", "Memory Pruning",
     "LRU eviction and 90-day TTL on agent context cache entries. Prevents unbounded memory growth in long-running deployments.",
     "app/services/memory_manager.py"),
    ("P2", "Cost Control", "Budget Period Reset Cron",
     "Monthly cron job resets per-org token/spend counters so budget caps apply per billing period rather than accumulating indefinitely.",
     "app/tasks/budget_reset.py"),
    ("P3", "Cost Optimisation", "Model Cost Tiering",
     "Formatting tasks (Mermaid diagram generation, JSON reshaping) routed to claude-haiku-4-5 (~10–15× cheaper). Core analysis tasks retain enterprise-grade models. 20–35% cost reduction per execution.",
     "app/services/llm_config.py → LIGHT_TASK_MODEL"),
    ("P3", "Multi-Tenancy", "Per-Org Anthropic API Key",
     "Enterprise customers supply their own Anthropic API key. Injected via Python ContextVar scoped to the executing asyncio Task — no signature changes required. Cleared in finally block to prevent cross-request leakage.",
     "app/services/llm_router.py → _org_anthropic_key"),
    ("P3", "Security", "Content Filtering on Uploads",
     "Uploaded documents (.txt, .md, .json, .csv) and URL-fetched content scanned for adversarial payloads before indexing. Returns 400/422 on detection.",
     "nextjs-app/src/app/api/context/upload/route.ts"),
    ("P3", "Security", "Recursive Input Sanitization",
     "All user-controlled fields in agent execution paths sanitized at injection time. Prevents prompt injection through nested JSON structures in inputData.",
     "nextjs-app/src/app/api/agents/execute/route.ts"),
    ("P3", "Reliability", "Fixed LLM Router Circular Dependency",
     "Eliminated circular call chain: _call_anthropic → claude_client → llm_router → _call_anthropic (infinite recursion in production). Rewrote _call_anthropic to call Anthropic SDK directly.",
     "app/services/llm_router.py"),
    ("P4", "Data Quality", "P4 Hallucination Detection",
     "Post-processing layer catches: negative times (critical), flow efficiency > 100% (critical), lead_time < PT+WT (critical), suspiciously round numbers (info), placeholder names (warning), identical risk scores (warning). Injects advisory _hallucination_flags dict; never blocks execution.",
     "app/services/hallucination_detector.py"),
    ("P4", "Experimentation", "A/B Testing Framework",
     "Deterministic SHA-256 variant assignment by execution_id (reproducible). Outcome metrics (steps discovered, confidence scores, hallucination flags) tracked per variant. Safe prompt experimentation without full rollout.",
     "app/services/ab_testing.py"),
    ("P4", "Data Quality", "Process Mining Engine",
     "Lightweight CSV event log ingestion (case_id, activity, timestamp). Computes activity stats, transition frequencies, bottleneck ranking (wait × frequency), and maps discovered timings to VSM steps via token similarity. Third independent evidence source alongside Jira and code signals.",
     "app/services/process_mining.py"),
    ("P4", "Performance", "Read Replica Database Pool",
     "Separate asyncpg pool for DATABASE_READ_URL. Analytics queries (benchmarks, metrics, process mining) offloaded to read replica. Falls back to primary if read URL not configured.",
     "app/core/database.py → ReadDatabasePool"),
    ("P4", "DevOps", "Kubernetes Manifests",
     "Production-ready K8s configs: HPA on agent-service (2–10 replicas) and Next.js (2–8 replicas), PostgreSQL StatefulSet with 50Gi PVC, Redis Deployment, nginx Ingress with TLS and SSE proxy-buffering off, rolling updates with maxUnavailable=0.",
     "k8s/ (9 manifest files)"),
]


def build_summary_docx(docx_path: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    run = p.add_run('TransformHub — P0 to P4 Enhancement Summary')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0C, 0x1C, 0x3A)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run('Production Hardening & Intelligence Quality Release  |  March 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    p.paragraph_format.space_after = Pt(12)

    # Overview paragraph
    p = doc.add_paragraph()
    run = p.add_run(
        'This document summarises all 29 enhancements shipped across five production hardening phases '
        '(P0 through P4). Phases are cumulative — each builds on the previous. Together they transform '
        'TransformHub from a functional prototype into an enterprise-grade, SOC2-aligned, multi-tenant '
        'platform with full audit trails, hallucination detection, process mining, and Kubernetes-native deployment.'
    )
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    # Phase summary table
    p = doc.add_heading(level=2)
    p.clear()
    run = p.add_run('Phase Overview')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    p.paragraph_format.space_after = Pt(6)

    phase_meta = [
        ('P0', 'Security & Enterprise Foundation', '8', 'Circuit breaker, SSO, RBAC, API keys, budget enforcement, rate limiting, prompt injection prevention'),
        ('P1', 'Data Quality & Auditability',      '5', 'Code signal extraction, Jira cycle time, manual overrides, schema validation, timing provenance'),
        ('P2', 'DevOps, Security & Scale',          '6', 'CI/CD, Row-Level Security, HNSW index, cross-org benchmarks, memory pruning, budget reset cron'),
        ('P3', 'Cost Optimisation & Security',      '5', 'Model tiering, per-org API keys, content filtering, input sanitization, circular dependency fix'),
        ('P4', 'Intelligence Quality & Operations', '5', 'Hallucination detection, A/B testing, process mining, read replica, Kubernetes manifests'),
    ]

    tbl = doc.add_table(rows=1 + len(phase_meta), cols=4)
    tbl.style = 'Table Grid'
    headers = ['Phase', 'Theme', 'Items', 'Key Capabilities']
    hdr_row = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shd(cell, '0C1C3A')

    for ri, (phase, theme, count, caps) in enumerate(phase_meta):
        row = tbl.rows[ri + 1]
        for j, text in enumerate([phase, theme, count, caps]):
            cell = row.cells[j]
            cell.text = text
            if cell.paragraphs[0].runs:
                r = cell.paragraphs[0].runs[0]
                r.font.size = Pt(9)
                if j == 0:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        if ri % 2 == 1:
            for j in range(4):
                _shd(row.cells[j], 'EFF6FF')

    doc.add_paragraph()

    # Detailed enhancements by phase
    current_phase = None
    phase_colours = {
        'P0': RGBColor(0x1E, 0x3A, 0x8A),
        'P1': RGBColor(0x07, 0x5D, 0x3A),
        'P2': RGBColor(0x78, 0x35, 0x0F),
        'P3': RGBColor(0x4C, 0x1D, 0x95),
        'P4': RGBColor(0x7F, 0x1D, 0x1D),
    }

    p = doc.add_heading(level=2)
    p.clear()
    run = p.add_run('Detailed Enhancement Log')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

    tbl2 = doc.add_table(rows=1 + len(ENHANCEMENTS), cols=5)
    tbl2.style = 'Table Grid'
    headers2 = ['Phase', 'Category', 'Enhancement', 'Description', 'File / Component']
    hdr_row2 = tbl2.rows[0]
    for j, h in enumerate(headers2):
        cell = hdr_row2.cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shd(cell, '1E3A8A')

    col_widths = [Inches(0.45), Inches(1.0), Inches(1.3), Inches(2.7), Inches(1.55)]
    for j, w in enumerate(col_widths):
        for row in tbl2.rows:
            row.cells[j].width = w

    for ri, (phase, cat, feat, desc, comp) in enumerate(ENHANCEMENTS):
        row = tbl2.rows[ri + 1]
        for j, text in enumerate([phase, cat, feat, desc, comp]):
            cell = row.cells[j]
            cell.text = text
            if cell.paragraphs[0].runs:
                r = cell.paragraphs[0].runs[0]
                r.font.size = Pt(8)
                if j == 0:
                    r.bold = True
                    r.font.color.rgb = phase_colours.get(phase, RGBColor(0, 0, 0))
                if j == 4:
                    r.font.name = 'Courier New'
                    r.font.size = Pt(7)
                    r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        if ri % 2 == 1:
            for j in range(5):
                _shd(row.cells[j], 'EFF6FF')

    doc.add_paragraph()

    # Closing notes
    p = doc.add_heading(level=2)
    p.clear()
    run = p.add_run('Test Coverage')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

    test_rows = [
        ('test_llm_router.py', '19', 'All routing, fallback, cost tiering, per-org key, token tracking'),
        ('test_hallucination_detector.py', '18', 'All severity levels: critical/warning/info; edge cases; dispatcher'),
        ('test_process_mining.py', '16', 'CSV parsing, activity stats, transitions, bottleneck ranking, VSM mapping'),
        ('test_ab_testing.py', '12', 'Deterministic assignment, outcome recording, experiment summary, edge cases'),
        ('Total', '65', ''),
    ]

    tbl3 = doc.add_table(rows=1 + len(test_rows), cols=3)
    tbl3.style = 'Table Grid'
    for j, h in enumerate(['Test File', 'Tests', 'Coverage']):
        cell = tbl3.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shd(cell, '1E3A8A')

    for ri, (f, n, cov) in enumerate(test_rows):
        row = tbl3.rows[ri + 1]
        for j, text in enumerate([f, n, cov]):
            cell = row.cells[j]
            cell.text = text
            if cell.paragraphs[0].runs:
                r = cell.paragraphs[0].runs[0]
                r.font.size = Pt(9)
                if j == 0: r.font.name = 'Courier New'; r.font.size = Pt(8)
                if ri == len(test_rows) - 1: r.bold = True
        if ri % 2 == 1:
            for j in range(3):
                _shd(row.cells[j], 'EFF6FF')

    doc.save(str(docx_path))
    print(f"  ✓  {docx_path.name}")


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    print("\nGenerating TransformHub P0–P4 Word documents...\n")

    jobs = [
        (DOCS_DIR / 'TransformHub-POV.md',                    DOCS_DIR / 'TransformHub-POV.docx'),
        (DOCS_DIR / 'TransformHub-CaseStudy-NationalBank.md', DOCS_DIR / 'TransformHub-CaseStudy-NationalBank.docx'),
        (DOCS_DIR / 'TransformHub-DemoGuide-P4.md',           DOCS_DIR / 'TransformHub-DemoGuide-P4.docx'),
    ]

    for md_path, docx_path in jobs:
        if not md_path.exists():
            print(f"  ✗  MISSING source: {md_path}")
            continue
        convert_md_to_docx(md_path, docx_path)

    build_summary_docx(DOCS_DIR / 'TransformHub-P0-P4-Enhancement-Summary.docx')

    print('\nDone. Files written to:', DOCS_DIR)
    print()
    for f in sorted(DOCS_DIR.glob('TransformHub-*.docx')):
        size_kb = f.stat().st_size // 1024
        print(f'  {f.name:55s}  {size_kb:>5} KB')
    print()


if __name__ == '__main__':
    main()
