"""Convert markdown file to Word document using python-docx."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_heading_style(para, level):
    colors = {
        1: RGBColor(0x1E, 0x3A, 0x8A),  # dark blue
        2: RGBColor(0x25, 0x63, 0xEB),  # medium blue
        3: RGBColor(0x3B, 0x82, 0xF6),  # lighter blue
        4: RGBColor(0x60, 0xA5, 0xFA),  # light blue
    }
    sizes = {1: 20, 2: 16, 3: 13, 4: 12}
    run = para.runs[0] if para.runs else para.add_run()
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))


def add_table_from_md(doc, lines, start_idx):
    """Parse markdown table and add as Word table."""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1

    if not table_lines:
        return start_idx

    # Parse headers
    headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
    # Skip separator row (line 1)
    data_rows = []
    for row_line in table_lines[2:]:
        cols = [c.strip() for c in row_line.split('|') if c.strip() != '']
        if cols:
            data_rows.append(cols)

    if not headers:
        return i

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for j, h in enumerate(headers[:num_cols]):
        cell = hdr_row.cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E3A8A')
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(data_rows):
        row = table.rows[ri + 1]
        for j in range(num_cols):
            cell = row.cells[j]
            text = row_data[j] if j < len(row_data) else ''
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            cell.text = text
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            # Alternate row shading
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EFF6FF')
                tcPr.append(shd)

    doc.add_paragraph()
    return i


def inline_format(para, text):
    """Add text to paragraph with inline bold/code formatting."""
    # Split on **bold** and `code` markers
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        else:
            run = para.add_run(part)
            run.font.size = Pt(10)


def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                # End code block — write as formatted paragraph
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_para.paragraph_format.left_indent = Inches(0.3)
                    code_para.paragraph_format.space_before = Pt(4)
                    code_para.paragraph_format.space_after = Pt(4)
                    for cl in code_lines:
                        run = code_para.add_run(cl + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                    # Add shading to code block
                    pPr = code_para._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'EFF6FF')
                    pPr.append(shd)
                in_code_block = False
                code_lines = []
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip mermaid diagrams (just add a placeholder note)
        if line.strip() == '```mermaid':
            in_code_block = True
            code_lines = ['[Diagram: see online documentation]']
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Table
        if line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            i = add_table_from_md(doc, lines, i)
            continue

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x0C, 0x1C, 0x3A)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if line.startswith('## '):
            p = doc.add_heading(level=2)
            p.clear()
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        if line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:].strip())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Blockquote
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            i += 1
            continue

        # Bullet list
        if re.match(r'^[-*+] ', line):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
            text = line[2:].strip()
            inline_format(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.3)
            text = re.sub(r'^\d+\. ', '', line).strip()
            inline_format(p, text)
            i += 1
            continue

        # Indented bullet (sub-list)
        if re.match(r'^  [-*+] ', line) or re.match(r'^    [-*+] ', line):
            p = doc.add_paragraph(style='List Bullet 2')
            p.paragraph_format.left_indent = Inches(0.6)
            text = re.sub(r'^[\s]+[-*+] ', '', line).strip()
            inline_format(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        inline_format(p, line.strip())
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    md_file = '/Users/125066/projects/TransformHub/docs/platform-docs/14-agent-orchestration.md'
    docx_file = '/Users/125066/projects/TransformHub/docs/TransformHub_Agent_Orchestration.docx'
    convert_md_to_docx(md_file, docx_file)
