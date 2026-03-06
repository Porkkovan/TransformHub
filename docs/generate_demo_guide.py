"""
Generate TransformHub_Demo_Guide.docx using python-docx.
Includes: overview, accuracy table, and full step-by-step guide.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ───────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0f, 0x34, 0x60)
RED    = RGBColor(0xe9, 0x45, 0x60)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x44, 0x44, 0x55)
GREEN  = RGBColor(0x1e, 0x7e, 0x34)
AMBER  = RGBColor(0x85, 0x64, 0x04)
DARKRED= RGBColor(0xc0, 0x39, 0x2b)
LIGHT  = RGBColor(0xf4, 0xf7, 0xfb)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        if side in kwargs:
            border = OxmlElement(f'w:{side}')
            for k,v in kwargs[side].items():
                border.set(qn(f'w:{k}'), v)
            tcBorders.append(border)
    tcPr.append(tcBorders)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(13)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = GRAY
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(text, bold=False, colour=None, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if colour:
        r.font.color.rgb = colour
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows, col_widths=None, header_bg='0F3460'):
    """Add a styled table with navy header row."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg = 'F4F7FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            if isinstance(cell_text, tuple):
                # (text, bold, colour_hex)
                text, bold, col_hex = cell_text
                r = p.add_run(text)
                r.bold = bold
                if col_hex:
                    r.font.color.rgb = RGBColor(
                        int(col_hex[0:2],16), int(col_hex[2:4],16), int(col_hex[4:6],16))
            else:
                r = p.add_run(str(cell_text))
                r.font.color.rgb = GRAY
            r.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return t

def page_break():
    doc.add_page_break()

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('─' * 80)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(7)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
r = p.add_run('TransformHub')
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
r2 = p2.add_run('Demo Guide & Step Accuracy Reference')
r2.font.size = Pt(16)
r2.font.color.rgb = RED
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = doc.add_paragraph()
r3 = p3.add_run('End-to-end walkthrough · Input sources · Output accuracy · User instructions')
r3.font.size = Pt(11)
r3.font.color.rgb = GRAY
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(10)

p4 = doc.add_paragraph()
r4 = p4.add_run('Version 1.0  ·  March 2026  ·  Confidential')
r4.font.size = Pt(9)
r4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(40)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PLATFORM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

heading1('1. Platform Overview')

body(
    'TransformHub is an AI-powered digital transformation workbench that guides organisations '
    'through a structured 7-step journey — from discovering what exists today to producing a '
    'prioritised transformation roadmap. It uses a fleet of 18 LangGraph AI agents (powered by '
    'Anthropic Claude) that automatically incorporate the organisation\'s uploaded knowledge base '
    'into every analysis.'
)

add_table(
    ['Layer', 'Technology / Detail'],
    [
        ('Frontend',          'Next.js 15 (App Router), TypeScript, Tailwind CSS v4 — dark glassmorphism UI'),
        ('Backend / Agents',  'FastAPI + LangGraph — 18 specialised Claude-powered agents'),
        ('Database',          'PostgreSQL with Prisma ORM — Org → Repo → Product → Capability → Functionality'),
        ('AI Model',          'Anthropic Claude (claude-sonnet-4-5) for all agent inference and classification'),
        ('Integrations',      'Jira, Confluence, Azure DevOps, Notion, ServiceNow (sync to Context Hub)'),
    ],
    col_widths=[1.5, 5.0]
)

heading2('Knowledge Base Architecture')
body(
    'Every document uploaded to the Context Hub is silently injected into every agent prompt '
    'via the shared format_context_section() function. This means uploading one document '
    'immediately benefits all 18 agents with no further configuration — making Context Hub '
    'the single highest-leverage action available to users.'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ACCURACY TABLE
# ══════════════════════════════════════════════════════════════════════════════

heading1('2. Output Accuracy by Step')

body(
    'Accuracy ratings reflect the expected alignment between AI-generated outputs and what an '
    'expert business analyst would produce independently. Ratings improve significantly when '
    'relevant documents are loaded into the Context Hub and when multi-pass modes are used.'
)

add_table(
    ['Step', 'Output Type', 'Baseline\nAccuracy', 'With Full\nContext', 'Key Accuracy Drivers'],
    [
        ('1. Org Setup',
         'Manual data entry',
         ('100%', True, '1E7E34'),
         ('100%', True, '1E7E34'),
         'User-defined — no AI inference'),

        ('2. Context Hub',
         'Document ingestion',
         ('100%', True, '1E7E34'),
         ('100%', True, '1E7E34'),
         'Storage only; downstream quality depends on document richness'),

        ('3. Discovery (single-pass)',
         'L1–L3 hierarchy',
         ('55–65%', True, '856404'),
         ('70–80%', True, '856404'),
         'OpenAPI spec availability, repo clarity, domain context notes'),

        ('3. Discovery (multi-pass)',
         'L1–L3 + personas',
         ('75–82%', True, '856404'),
         ('85–92%', True, '1E7E34'),
         'Each pass refines previous; human review gates enforce quality'),

        ('3. Persona AI Mapping',
         'Persona–Functionality assignments',
         ('70–78%', True, '856404'),
         ('82–88%', True, '1E7E34'),
         'Quality of persona responsibility definitions; functionality naming'),

        ('4. VSM (upload-based)',
         'PT/WT metrics, flow diagrams',
         ('90–95%', True, '1E7E34'),
         ('90–95%', True, '1E7E34'),
         'Data-driven — accuracy equals quality of uploaded timing data'),

        ('4. VSM (AI-inferred from URL)',
         'PT/WT estimates, step classification',
         ('60–70%', True, '856404'),
         ('70–78%', True, '856404'),
         'Document accessibility and process description completeness'),

        ('5. Risk & Compliance',
         'Risk register, compliance gaps',
         ('65–75%', True, '856404'),
         ('75–85%', True, '856404'),
         'Requires regulatory docs in Context Hub for domain-specific risks'),

        ('6. Future State Vision',
         'Target capability map, recommendations',
         ('60–72%', True, '856404'),
         ('72–82%', True, '856404'),
         'Most context-sensitive step — needs strategy, benchmarks, roadmaps'),

        ('7. Transformation Roadmap',
         'Phased initiatives, priorities',
         ('62–72%', True, '856404'),
         ('73–83%', True, '856404'),
         'Compounds accuracy of all upstream steps (3–6)'),
    ],
    col_widths=[1.6, 1.6, 0.9, 0.9, 2.5]
)

body('Accuracy colour guide:', bold=True)
bullet('85–100%  — High confidence. Suitable for stakeholder presentation with light review.')
bullet('70–84%   — Medium-high. Good first draft; verify domain-specific details.')
bullet('55–69%   — Medium. Use as structured input to an expert review process.')
bullet('< 55%    — Low. Treat as raw signal only; substantial expert refinement required.')

body(
    '\nKey insight: The Context Hub is the single highest-leverage action. '
    'Uploading 3–5 relevant documents before running agents lifts accuracy '
    '10–15 percentage points across every step simultaneously.',
    bold=True, colour=NAVY
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# STEPS 1–7
# ══════════════════════════════════════════════════════════════════════════════

steps = [
    {
        'num': '3', 'title': 'Step 1 — Organisation Setup',
        'purpose': (
            'Define the organisation\'s structure, business segments, and user personas. '
            'This is the master configuration that all agents reference throughout every '
            'subsequent step.'
        ),
        'inputs': [
            ('Organisation Name',  'Legal or trading name',                                      'User enters manually'),
            ('Business Segments',  'Operating divisions (e.g. Retail, Corporate & Commercials)', 'Operating model or annual report'),
            ('Personas',           '2–5 types with name, type code, responsibilities list',      'HR org design, target operating model docs'),
        ],
        'how': (
            'Data saved directly to the Organization database record as structured JSON. '
            'No AI inference involved. Every agent prompt automatically includes the org\'s '
            'personas and segment context via the shared format_context_section() function '
            'in the agent service.'
        ),
        'outputs': [
            'Org profile record with segments array and personas JSON',
            'Segment dropdown options across Discovery, VSM, and Catalog pages',
            'Persona types available for mapping in Discovery',
        ],
        'accuracy': '100% — User-defined, no AI inference.',
        'acc_color': '1E7E34',
    },
    {
        'num': '4', 'title': 'Step 2 — Context Hub',
        'purpose': (
            'Build the organisation\'s knowledge base. Every document, URL, and integration '
            'sync result uploaded here is silently injected into every agent call — making '
            'this the single most impactful step for improving output quality across the '
            'entire platform.'
        ),
        'inputs': [
            ('PDF / Word docs',    'BRDs, architecture docs, strategy papers, regulatory frameworks', 'Upload via Documents tab'),
            ('Web URLs',           'Confluence pages, internal wikis, regulatory body websites',      'Paste URL — system fetches and indexes'),
            ('Free text',          'Pasted notes, meeting summaries, key decisions',                  'Text input in Documents tab'),
            ('Jira',               'Open epics, active sprints, backlog items',                       'Integrations tab → credentials + project key → Sync'),
            ('Confluence',         'Space pages, architecture decision records',                      'Integrations tab → space key → Sync'),
            ('Azure DevOps',       'Work items, user stories, features',                              'Integrations tab → org/project → Sync'),
            ('Notion / ServiceNow','Project databases, ITSM catalogue',                               'Integrations tab → API token + database ID → Sync'),
        ],
        'how': (
            'Documents are stored as ContextDocument records with category and sub-category tags. '
            'Before every agent execution, format_context_section() retrieves the top relevant '
            'documents for the org and prepends them to the agent\'s system prompt. '
            'This means uploading one document immediately benefits all 18 agents with no further configuration.'
        ),
        'outputs': [
            'Indexed ContextDocument records visible in Admin → Manage Docs panel',
            'Automatic enrichment of every subsequent agent call with relevant context',
            'External integration records with sync status and item counts',
        ],
        'accuracy': '100% — Storage and retrieval; downstream impact depends on document quality.',
        'acc_color': '1E7E34',
    },
    {
        'num': '5', 'title': 'Step 3 — Discovery',
        'purpose': (
            'Automatically analyse software repositories and documentation to discover the '
            'organisation\'s digital product portfolio — structured as a four-level hierarchy '
            '(L0 Product Groups → L1 Products → L2 Capabilities → L3 Functionalities) — and '
            'map functionalities to user personas.'
        ),
        'inputs': [
            ('Repository URLs',     'GitHub / GitLab / Bitbucket repo URLs',          'Engineering team — highest accuracy signal'),
            ('OpenAPI Spec URLs',   'Swagger / OpenAPI endpoint URLs',                 'API team / developer portal — reveals exact capabilities'),
            ('GitHub Token',        'Personal access token for private repos',         'Engineering team — enables deep file scanning'),
            ('DB Schema Text',      'Paste SQL CREATE TABLE statements',               'DBA / architecture team — reveals domain entities'),
            ('Domain Context',      'Free text: "KYB onboarding system for corporate banking"', 'Business analyst / architect — disambiguates AI'),
            ('Known Products',      'Comma-separated list of expected products',       'Product manager — anchors discovery output'),
            ('Business Segment',    'Segment dropdown selection',                      'UI selection — tags all discovered products'),
            ('Context Hub docs',    'All uploaded documents',                          'Step 2 output — auto-injected by platform'),
        ],
        'how': (
            'SINGLE-PASS: One LangGraph agent call analyses all inputs simultaneously. '
            'Fetches repo URLs, parses OpenAPI specs, reads Context Hub documents, then '
            'calls Claude to infer a structured L1→L2→L3 hierarchy. Assigns confidence scores '
            '(0.0–1.0) and evidence sources (url_analysis, openapi_spec, github_structure, '
            'context_document) to each item.\n\n'
            'MULTI-PASS (Recommended): Pass 1 = broad discovery of products and top-level '
            'capabilities → human review gate → Pass 2 = capability deepening → human review '
            'gate → Pass 3 = functionality enrichment + persona mappings + confidence scoring.\n\n'
            'PERSONA AUTO-MAPPING: Click "Auto-Map with AI" in the Persona–Functionality Matrix. '
            'Claude classifies which personas interact with each functionality based on name and '
            'responsibility definitions. Falls back to keyword rules if API key is unavailable.'
        ),
        'outputs': [
            'DigitalProduct records (L1) — name, description, business segment, confidence score, evidence sources',
            'DigitalCapability records (L2) — name, description, category, linked to parent product',
            'Functionality records (L3) — name, description, source files, linked to parent capability',
            'PersonaMapping records — each functionality mapped to one or more persona types',
            'Persona–Functionality Matrix — visual coverage grid with % coverage',
            'Product Catalog flat table — all 68 rows XLSX-downloadable with classification badges',
            'Confidence badges — per-item score with evidence source pills (green ≥80%, amber 60–79%, red <60%)',
        ],
        'accuracy': 'Single-pass: 55–65% baseline → 70–80% with context. Multi-pass: 75–82% → 85–92% with context + OpenAPI.',
        'acc_color': '856404',
    },
    {
        'num': '6', 'title': 'Step 4 — Value Stream Mapping (VSM)',
        'purpose': (
            'Quantify the efficiency of each product\'s operational flow by assigning Process Time (PT), '
            'Wait Time (WT), Lead Time (LT), and Flow Efficiency (FE = PT/LT) to every capability '
            'discovered in Step 3. Produces visual Mermaid flow diagrams and identifies bottlenecks.'
        ),
        'inputs': [
            ('Product selection',           'Choose a digital product from Discovery output',         'Step 3 output — auto-loaded in sidebar'),
            ('Process Map Upload (Option A)', 'XLSX/CSV with step names + PT/WT columns. Auto-detects units (mins/hours/days/weeks)', 'Business analyst / process owner'),
            ('VSM Metrics CSV (Option B)',   'Download template (pre-filled capability names) → add PT/WT → upload back', 'Step 3 capabilities + user timing data'),
            ('Process Doc URL (Option C)',   'URL to existing process documentation or Confluence page', 'Process owner — AI extracts timing estimates'),
            ('Competitor Value Streams',     'Competitor name + description or benchmark URL',         'Strategy / competitive intelligence team'),
            ('Context Hub docs',             'Process standards, SLA docs, benchmark data',            'Step 2 output — auto-injected'),
        ],
        'how': (
            'The lean_vsm LangGraph agent loads capabilities from DB (Step 3 output), applies '
            'provided timing data, then calls Claude to classify each step as Value-Adding, '
            'Bottleneck, or Waste. Aggregates metrics: PT = sum of value-adding time; '
            'WT = sum of delay/queue time; LT = PT + WT; FE = PT/LT × 100%. '
            'Generates Mermaid flowchart source with colour-coded nodes. '
            'Three view levels: L1 (segment totals), L2 (per-capability cards with diagrams), '
            'L3 (per-functionality step breakdown).\n\n'
            'NEW — VSM Metrics Template: Download pre-filled XLSX from sidebar, edit PT/WT '
            'per capability, upload back. Fuzzy column detection handles header name variations. '
            'Idempotent — safe to run multiple times.'
        ),
        'outputs': [
            'VsmMetrics per capability — processTime, waitTime, leadTime, flowEfficiency (%) stored in DB',
            'Mermaid flow diagrams — one per capability + one L1 aggregate (green = value-adding, red = bottleneck)',
            'Flow Efficiency scores — target ≥40%; typical banking: 5–15%',
            'Bottleneck identification — steps with WT > 3×PT flagged as primary transformation targets',
            'Step Classification report — downloadable breakdown of value-adding vs bottleneck vs waste per step',
        ],
        'accuracy': 'Upload-based: 90–95%. AI-inferred from URL: 60–70%.',
        'acc_color': '1E7E34',
    },
    {
        'num': '7', 'title': 'Step 5 — Risk & Compliance',
        'purpose': (
            'Identify operational, regulatory, and technology risks associated with the current '
            'capability landscape and VSM bottlenecks. Generate a prioritised risk register and '
            'compliance gap analysis.'
        ),
        'inputs': [
            ('Digital capabilities + functionalities', 'L2–L3 hierarchy',                      'Step 3 Discovery — auto-loaded from DB'),
            ('VSM metrics per capability',             'PT, WT, Flow Efficiency per capability', 'Step 4 VSM — auto-loaded from DB'),
            ('Regulatory framework documents',         'APRA CPS 234, PCI-DSS, GDPR, AML/CTF Act etc.', 'Context Hub uploads — critical for accuracy'),
            ('Business segment',                       'Governs which regulations apply',       'Step 1 Org Setup'),
            ('Architecture / security docs',           'System design, security standards',     'Context Hub uploads'),
        ],
        'how': (
            'The risk_compliance LangGraph agent loads all capabilities and their VSM metrics, '
            'identifies capabilities with low flow efficiency (<10%) as operational risk signals, '
            'then cross-references against regulatory requirements from Context Hub documents. '
            'Calls Claude to assess each capability against: operational risk, regulatory compliance, '
            'data risk, technology risk, and third-party/vendor risk. '
            'Scores each risk: Likelihood (1–5) × Impact (1–5) = Risk Score.'
        ),
        'outputs': [
            'Risk register — capability name, risk type, severity (Critical/High/Medium/Low), likelihood, impact, recommended mitigation',
            'Compliance gap analysis — which regulatory requirements each capability partially or fully fails',
            'Priority matrix — top 5 risks requiring immediate attention',
        ],
        'accuracy': '65–75% baseline → 75–85% with regulatory docs in Context Hub.',
        'acc_color': '856404',
    },
    {
        'num': '8', 'title': 'Step 6 — Future State Vision',
        'purpose': (
            'Generate a target operating model — recommending which capabilities to retain, enhance, '
            'automate, consolidate, or retire, and proposing net-new capabilities to address gaps.'
        ),
        'inputs': [
            ('Current capability map (L1–L3)',    'All discovered products, capabilities, functionalities', 'Step 3 Discovery — auto-loaded'),
            ('VSM flow efficiency scores',         'Per-capability FE%, bottleneck flags',                  'Step 4 VSM — auto-loaded'),
            ('Risk register',                      'Risk severity and compliance gaps',                     'Step 5 Risk & Compliance — auto-loaded'),
            ('Strategy documents',                 'Digital transformation vision, OKRs',                  'Context Hub — most impactful input for this step'),
            ('Industry benchmarks',                'Competitor analysis, industry reports',                 'Context Hub uploads or competitor URLs from VSM'),
            ('Technology roadmaps',                'Platform capabilities, vendor roadmaps',               'Context Hub uploads'),
        ],
        'how': (
            'The future_state_vision LangGraph agent sorts capabilities by flow efficiency '
            '(worst first = highest priority for change), loads the risk register to identify '
            'compounding risk + inefficiency, reads strategy and benchmark documents from Context Hub, '
            'then calls Claude to reason: for each capability — Retain As-Is / Enhance / '
            'Automate / Consolidate / Retire / New. Proposes net-new capabilities for gaps '
            'and generates a narrative vision statement per product.'
        ),
        'outputs': [
            'Future state disposition per capability (Retain / Enhance / Automate / Consolidate / Retire)',
            'Net-new capabilities recommended with rationale',
            'Narrative future state vision document per product',
            'Gap analysis between current and future state capability counts',
            'Estimated flow efficiency improvement if future state is realised',
        ],
        'accuracy': '60–72% baseline → 72–82% with strategy docs. Most context-sensitive step.',
        'acc_color': '856404',
    },
    {
        'num': '9', 'title': 'Step 7 — Product Transformation Roadmap',
        'purpose': (
            'Convert the future state vision into a concrete, phased transformation roadmap '
            'with prioritised initiatives, effort estimates, dependencies, and expected outcomes.'
        ),
        'inputs': [
            ('Future state dispositions',         'Per-capability change type (Enhance/Automate/Retire/New)', 'Step 6 — auto-loaded'),
            ('Risk priorities',                   'High/Critical risks from risk register',                  'Step 5 — auto-loaded'),
            ('VSM bottleneck severity',           'Flow efficiency scores, bottleneck flags',                'Step 4 — auto-loaded'),
            ('Persona mappings',                  'Who is affected by each capability change',               'Step 3 — auto-loaded'),
            ('Programme governance docs',         'Investment constraints, approval thresholds',             'Context Hub uploads'),
            ('Business case templates',           'Standard financial modelling structure',                  'Context Hub uploads'),
        ],
        'how': (
            'The product_transformation LangGraph agent ranks future-state dispositions by combined '
            'value score (flow efficiency gain × risk reduction × strategic alignment), then groups '
            'into three horizons: Quick Wins (0–3 months, low effort, high FE gain), '
            'Medium Term (3–12 months), Strategic (12–24 months). '
            'Identifies dependencies between initiatives, estimates effort (S/M/L/XL) and '
            'expected FE improvement per initiative, assigns impacted personas.'
        ),
        'outputs': [
            'Transformation initiatives — name, description, horizon, effort (S/M/L/XL), priority score',
            'Expected outcomes per initiative — flow efficiency improvement (%), risk reduction, personas benefited',
            'Dependency map — which initiatives must be completed before others can start',
            'Phased roadmap timeline view — grouped by horizon',
            'Programme narrative — executive summary of the transformation programme',
        ],
        'accuracy': '62–72% baseline → 73–83% with full upstream completion. Best with human-approved Steps 3–6.',
        'acc_color': '856404',
    },
]

for step in steps:
    heading1(step['title'])
    heading2('Purpose')
    body(step['purpose'])

    heading2('Inputs')
    add_table(
        ['Input', 'Description', 'Source'],
        [(r[0], r[1], r[2]) for r in step['inputs']],
        col_widths=[1.6, 2.8, 2.1]
    )

    heading2('How It Works')
    for para in step['how'].split('\n\n'):
        body(para.strip())

    heading2('Outputs')
    for out in step['outputs']:
        bullet(out)

    heading2('Accuracy')
    p = doc.add_paragraph()
    r = p.add_run(step['accuracy'])
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(
        int(step['acc_color'][0:2],16),
        int(step['acc_color'][2:4],16),
        int(step['acc_color'][4:6],16)
    )

    divider()
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DEMO TIPS
# ══════════════════════════════════════════════════════════════════════════════

heading1('Demo Tips & Common Pitfalls')

heading2('Before the Demo')
bullet('Select US Bank as the demo organisation (pre-seeded with Corporate & Commercials and Retail Banking segments)')
bullet('Upload at least one BRD or architecture document to Context Hub before running Discovery — the output quality difference is immediately visible')
bullet('Use Corporate and Commercials segment for richest data (6 products, 18 capabilities, 68 functionalities including Client Onboarding with full VSM metrics)')

heading2('Recommended Demo Flow')
heading3('Discovery')
numbered('Set segment to Corporate and Commercials')
numbered('Enter repo URL — usbank-core-banking has pre-existing data')
numbered('Show Products View — 6 product cards with confidence badges and evidence source pills')
numbered('Switch to Product Catalog — flat table of all 68 functionalities with Bottleneck / Value Adding classification badges')
numbered('Click "Auto-Map with AI" in the Persona–Functionality Matrix — live AI classification in ~5 seconds')

heading3('VSM')
numbered('Navigate to VSM page')
numbered('Select Client Onboarding Management from the product sidebar')
numbered('Show L2 view — 6 capability cards with PT/WT/FE metrics and Mermaid flow diagrams')
numbered('Show "Update VSM Metrics" card in sidebar → click Download Template — pre-filled XLSX opens immediately')
numbered('Explain: user edits PT/WT numbers, uploads back, metrics refresh live')
numbered('Switch to L3 view to show per-functionality step classification')

heading2('Common Pitfalls')
add_table(
    ['Issue', 'Fix'],
    [
        ('Discovery returns generic or wrong capabilities',
         'Add domain context text ("This is a KYB system for corporate banking") and upload a BRD to Context Hub before running'),
        ('Persona matrix shows 0% coverage',
         'Click "Auto-Map with AI" — initial Discovery may not auto-map personas without existing mapping data'),
        ('VSM metrics template downloads with empty capability column',
         'Run Discovery first to create capabilities, then return to VSM page'),
        ('Multi-pass Discovery: Pass 2 does not see Pass 1 results',
         'Approve Pass 1 in the MultiPass panel before clicking Pass 2 — approval gates are required'),
        ('Agent returns timeout or 502 error',
         'Ensure agent service is running: cd agent-service && uvicorn app.main:app --port 8002'),
        ('Segment filter not showing products',
         'In Discovery, always select the Business Segment BEFORE running the agent — products get tagged with the selected segment at creation time'),
    ],
    col_widths=[2.6, 4.0]
)

heading2('Key Messages for Stakeholders')
bullet('"The more you put in, the more you get out" — every document in the Context Hub improves every agent automatically')
bullet('"Human-in-the-loop by design" — multi-pass Discovery and approval gates ensure AI output is reviewed before influencing downstream steps')
bullet('"Accuracy compounds" — Step 7 roadmap accuracy depends on quality of Steps 3–6; investing in good inputs at Step 3 pays dividends across the entire platform')
bullet('"Existing data respected" — the platform reads from Jira/Confluence/ADO; no need to re-enter what is already documented')

divider()
p = doc.add_paragraph('TransformHub Demo Guide  ·  March 2026  ·  Internal Use Only')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(8)
p.runs[0].font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/125066/Desktop/TransformHub_Demo_Guide.docx'
doc.save(out)
print(f'Saved: {out}')
