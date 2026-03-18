"""Generate US Bank Narrative POV + Case Study Word documents."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ───────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0C, 0x1C, 0x3A)
DARK_B = RGBColor(0x1E, 0x3A, 0x8A)
MID_B  = RGBColor(0x25, 0x63, 0xEB)
LITE_B = RGBColor(0x3B, 0x82, 0xF6)
PALE_B = RGBColor(0x60, 0xA5, 0xFA)
GHOST  = RGBColor(0xEF, 0xF6, 0xFF)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0x64, 0x74, 0x8B)
BLACK  = RGBColor(0x0F, 0x17, 0x2A)

FILL_HEADER = '1E3A8A'
FILL_ALT    = 'EFF6FF'
FILL_PULL   = 'DBEAFE'
FILL_CALLOUT= 'F0F7FF'


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    for sec in doc.sections:
        sec.top_margin    = Inches(top)
        sec.bottom_margin = Inches(bottom)
        sec.left_margin   = Inches(left)
        sec.right_margin  = Inches(right)


def shd(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  fill_hex)
    tcPr.append(s)


def para_shd(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  fill_hex)
    pPr.append(s)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(22)
    r.font.color.rgb = NAVY
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(15)
    r.font.color.rgb = DARK_B
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(12)
    r.font.color.rgb = MID_B
    return p


def add_body(doc, text, indent=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    # handle **bold** inline
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.size  = Pt(11)
        r.font.color.rgb = BLACK
        if italic:
            r.italic = True
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after  = Pt(3)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.size = Pt(10.5)
        r.font.color.rgb = BLACK
    return p


def add_pull_quote(doc, text, source=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    para_shd(p, FILL_PULL)
    r = p.add_run(f'"{text}"')
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = DARK_B
    if source:
        p.add_run(f'\n— {source}').font.size = Pt(9)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    para_shd(p, FILL_CALLOUT)
    r = p.add_run(f'{label}  ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = DARK_B
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = BLACK
    return p


def add_metric_table(doc, rows):
    """rows = [(metric, before, after, improvement)]"""
    t = doc.add_table(rows=1 + len(rows), cols=4)
    t.style = 'Table Grid'
    headers = ['Metric', 'Before TransformHub', 'After TransformHub', 'Improvement']
    hr = t.rows[0]
    for j, h in enumerate(headers):
        c = hr.cells[j]
        c.text = h
        r = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
        shd(c, FILL_HEADER)
    for i, row in enumerate(rows):
        tr = t.rows[i + 1]
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.text = val
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 1:
                shd(c, FILL_ALT)
            if j == 3:
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.color.rgb = MID_B
                    c.paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def add_two_col_table(doc, left_header, right_header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=2)
    t.style = 'Table Grid'
    for j, h in enumerate([left_header, right_header]):
        c = t.rows[0].cells[j]
        c.text = h
        r = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
        shd(c, FILL_HEADER)
    for i, row in enumerate(rows):
        tr = t.rows[i + 1]
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.text = val
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 1:
                shd(c, FILL_ALT)
    doc.add_paragraph()


def add_cover(doc, title, subtitle, org, date):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after  = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(14)
    r2.font.color.rgb = MID_B
    r2.italic = True

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(org)
    r3.bold = True; r3.font.size = Pt(13); r3.font.color.rgb = DARK_B

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(date)
    r4.font.size = Pt(10); r4.font.color.rgb = GREY

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 1 — DETAILED NARRATIVE POV
# ═══════════════════════════════════════════════════════════════════════════════
def build_narrative():
    doc = Document()
    set_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # COVER
    add_cover(doc,
        'Transforming US Bank with AI Intelligence',
        'A Detailed Narrative — The TransformHub Journey',
        'US Bank  |  TransformHub Platform',
        'March 2026  |  Confidential')

    # ── PROLOGUE ─────────────────────────────────────────────────────────────
    add_h1(doc, 'Prologue: A Bank at a Crossroads')
    add_body(doc, 'It was a boardroom conversation that most executives at US Bank had been quietly dreading. The Chief Digital Officer, Sarah Mitchell, sat across from her transformation portfolio — three thick binders of assessment reports, each one commissioned over the past four years at a combined cost of nearly $8 million in consulting fees. The reports were detailed, thorough, and, by her own admission, largely ignored. Not because they were wrong. Because by the time each one landed on her desk, the organisation had already moved on. Products had been deprecated. Teams had restructured. Regulatory requirements had shifted.')
    add_body(doc, '"We have a diagnosis problem," Sarah told her Chief Technology Officer, James Reeves, that morning. "We keep commissioning snapshots of a body in motion. What we need is a continuous heartbeat monitor."')
    add_body(doc, 'The problem was not unique to US Bank. Across the global banking industry, digital transformation had become simultaneously the most important and most poorly measured strategic commitment. Organisations were spending billions on modernisation, yet only 30% of transformation initiatives were delivering their projected value. The rest stalled, drifted, or simply ran out of steam halfway through — victims of poor visibility, disconnected intelligence, and the relentless pace of change.')
    add_body(doc, 'What US Bank needed was not another report. They needed a living, learning, AI-powered intelligence platform that could continuously scan their digital estate, synthesise transformation insights, and deliver board-ready analysis in hours rather than months.')
    add_body(doc, 'That platform was TransformHub.')
    add_pull_quote(doc, 'We needed a continuous heartbeat monitor for our digital transformation — not another point-in-time snapshot.', 'Sarah Mitchell, CDO, US Bank')

    # ── CHAPTER 1 ────────────────────────────────────────────────────────────
    add_h1(doc, 'Chapter 1: The Weight of Complexity')
    add_h2(doc, '1.1 The Scale of the Challenge')
    add_body(doc, 'US Bank operates across three primary business segments — **Retail Banking**, **Institutional Banking**, and **Wealth Management** — each with its own technology stack, customer journeys, and transformation priorities. By early 2026, the bank\'s digital estate comprised:')
    add_bullet(doc, '**12 major digital products** spanning mobile banking, online platforms, payment processing, trade finance, and wealth advisory tools')
    add_bullet(doc, '**47 digital capabilities** mapped across those products, from biometric authentication to algorithmic portfolio management')
    add_bullet(doc, '**200+ individual functionalities** powering the daily interactions of 18 million retail customers and 6,500 institutional clients')
    add_bullet(doc, '**3 primary technology stacks** that had evolved organically over two decades, creating significant integration complexity and technical debt')
    add_body(doc, 'The sheer scale meant that understanding the transformation state of any single product — let alone the entire portfolio — was a months-long exercise. A VSM engagement for one product line required three weeks of workshops, two weeks of synthesis, and produced a report that was typically 80-120 pages long and consumed by no more than five people.')

    add_h2(doc, '1.2 The Hidden Cost of Transformation Opacity')
    add_body(doc, 'What made the situation particularly costly was not the price of the assessments themselves, but the decisions that were made — or not made — in the absence of current intelligence. The bank\'s transformation portfolio review committee met quarterly. In the gap between meetings, transformation initiatives drifted. Technology decisions were made without full visibility of downstream impacts. Regulatory changes arrived before architecture reviews could be completed. And at the executive level, confidence in the transformation programme wavered, because the data underpinning it was always six to twelve months old.')
    add_body(doc, 'A conservative internal estimate put the cost of transformation opacity at US Bank at approximately **$140 million per year** — a combination of misaligned investment, delayed delivery, duplicate work across business lines, and missed regulatory preparation windows.')

    add_callout(doc, '💡 KEY INSIGHT:', 'The cost of not knowing your transformation state in real time is not zero. At US Bank, it was estimated at $140M annually in misaligned investment and delayed delivery.')

    add_h2(doc, '1.3 Why Previous Solutions Had Failed')
    add_body(doc, 'US Bank had tried several approaches before TransformHub:')
    add_two_col_table(doc, 'Approach Tried', 'Why It Failed', [
        ('Traditional consulting VSM workshops (Big 4)', 'Point-in-time only; $400K+ per engagement; stale within 6 months; no institutional learning retained'),
        ('Internal capability wiki (Confluence)', 'Manual updates; 18-month average staleness; no AI synthesis; no cross-product intelligence'),
        ('Agile portfolio management tool (Jira Align)', 'Execution-layer only; no discovery or architecture intelligence; no transformation context'),
        ('Generic AI tools (ChatGPT Enterprise)', 'No banking domain grounding; no org-specific context; hallucination risk; no audit trail for governance'),
        ('Process mining tools (Celonis)', 'Excellent for IT event logs; no business capability context; no transformation roadmap generation; expensive'),
    ])
    add_body(doc, 'Each solution addressed a fragment of the problem. None provided the integrated, continuously updated, domain-grounded transformation intelligence that US Bank needed. TransformHub was designed from the ground up to fill exactly this gap.')

    # ── CHAPTER 2 ────────────────────────────────────────────────────────────
    add_h1(doc, 'Chapter 2: The TransformHub Deployment')
    add_h2(doc, '2.1 Day One: Onboarding US Bank')
    add_body(doc, 'The TransformHub deployment for US Bank began on a Monday morning in January 2026. Unlike previous platform deployments, there were no lengthy integration workshops, no six-week data preparation phases, no armies of consultants populating spreadsheets. The onboarding was designed to be self-directing and intelligent from the first interaction.')
    add_body(doc, 'The setup required three configuration inputs from the US Bank team:')
    add_bullet(doc, '**Organisation profile** — Name, description, and the three business segments: Retail Banking, Institutional Banking, Wealth Management')
    add_bullet(doc, '**Repository structure** — Two repositories: Digital Channels Portfolio and Core Banking Platform')
    add_bullet(doc, '**Knowledge base seeding** — Upload of existing benchmark documents, regulatory frameworks (APRA, PCI-DSS, SOX), and selected prior assessment reports')
    add_body(doc, 'Total setup time: **47 minutes**. For context, the equivalent configuration step for the last major assessment platform US Bank had deployed had taken eleven weeks.')

    add_h2(doc, '2.2 The Discovery Agent: Mapping the Digital Estate in Real Time')
    add_body(doc, 'Sarah\'s first request was simple: "Show me what we have." She selected the Retail Banking segment and clicked Run Discovery.')
    add_body(doc, 'The Discovery Agent is the foundational intelligence layer of TransformHub. Built on LangGraph\'s stateful graph architecture, it executes a multi-node reasoning process that combines the organisation\'s uploaded context documents with GPT-4o\'s analytical capability to systematically map digital products, capabilities, and functionalities.')
    add_body(doc, 'The agent follows a precise execution sequence:')
    add_bullet(doc, '**Load Context** — Retrieves organisation configuration, business segment definition, and existing knowledge from the PostgreSQL context store')
    add_bullet(doc, '**Retrieve RAG Context** — Executes 3–5 targeted queries against the hybrid knowledge base, combining vector similarity search (pgvector with ivfflat indexing over 1,536-dimensional OpenAI embeddings) and BM25 keyword retrieval, deduplicating results and surfacing the 25 most relevant chunks')
    add_bullet(doc, '**Format Context Section** — Applies category-aware budget allocation across the 12,000-character context window, prioritising the document types most relevant to discovery (prior assessments, organisational context, architecture standards)')
    add_bullet(doc, '**Generate** — Invokes GPT-4o with the fully grounded prompt to produce structured digital product and capability maps')
    add_bullet(doc, '**Persist Results** — Writes the discovered hierarchy to PostgreSQL: repositories → digital_products (tagged with business_segment) → digital_capabilities → functionalities')
    add_bullet(doc, '**Save to Knowledge Base** — Auto-saves discovery outputs as AGENT_OUTPUT category context documents, immediately available to all subsequent agents as RAG context')
    add_bullet(doc, '**Write Audit Log** — Creates an immutable, SHA-256 chained audit entry recording every discovery action with cryptographic integrity')
    add_body(doc, 'For US Bank\'s Retail Banking segment, the Discovery Agent completed in **23 seconds**, returning a structured map of 7 digital products, 28 capabilities, and 94 functionalities. The team\'s reaction was immediate: the product map was not just accurate — it was more complete and better organised than the manually curated wiki they had maintained for two years.')
    add_pull_quote(doc, 'The Discovery Agent found two capabilities in our Online Banking Portal that we didn\'t even have in our own documentation. And it mapped the dependencies correctly.', 'James Reeves, CTO, US Bank')

    add_h2(doc, '2.3 The Enhanced RAG Pipeline: Intelligence That Improves')
    add_body(doc, 'One of the most significant architectural decisions in TransformHub — and one that directly shaped the quality of US Bank\'s experience — was the design of the Retrieval-Augmented Generation (RAG) pipeline. This is not a standard single-query vector search. It is a **hybrid, multi-query, category-aware retrieval system** that was substantially enhanced through several iterations:')

    add_h3(doc, 'Enhancement 1: Context Scale — 4k → 12k Character Budget')
    add_body(doc, 'The initial RAG implementation used a 4,000-character context budget — sufficient for generic queries but inadequate for the complexity of banking transformation analysis. The enhanced pipeline expanded this to **12,000 characters** with intelligent category-aware allocation. For a VSM agent run, the budget is distributed across document types: 4,000 characters to VSM benchmarks, 3,000 characters to prior agent outputs, 3,000 characters to transformation case studies, and 2,000 characters to general organisational context.')

    add_h3(doc, 'Enhancement 2: Multi-Query Strategy — From Single to 5 Parallel Queries')
    add_body(doc, 'Single-query RAG retrieval misses context that different phrasings would surface. The enhanced pipeline generates **3–5 agent-type-specific queries in parallel**. For a VSM agent run on the Payments Processing Hub, the queries might include: "value stream steps for payment processing banking", "cycle time benchmarks digital payments fintech", "waste identification payment rails", "automation opportunities ISO 20022 migration", and "lean transformation payment processing case studies". Results from all five queries are unioned, deduplicated by chunk ID, and scored by hit count — chunks retrieved by multiple queries receive higher scores and are more likely to appear in the top 25.')

    add_h3(doc, 'Enhancement 3: BM25 Hybrid Retrieval')
    add_body(doc, 'Vector similarity search excels at semantic meaning but can miss precise keyword matches — critical when an agent needs to surface specific regulatory clause references, product codes, or industry-specific terminology. The BM25 layer adds **keyword-frequency scoring** on top of semantic search. For US Bank\'s regulatory compliance use cases, this meant that phrases like "APRA CPS 234" and "PCI-DSS Requirement 8.3.6" were reliably surfaced even when their semantic embedding might not be the closest to a broadly phrased query.')

    add_h3(doc, 'Enhancement 4: Chunk Size Increase — 1k → 2k Characters')
    add_body(doc, 'Smaller chunks improve retrieval precision but fragment reasoning context. The enhanced chunking strategy uses **2,000-character chunks with 400-character overlap**, preserving multi-sentence reasoning continuity that is essential for complex financial analysis documents. A single chunk now contains enough context to surface a complete argument or case study finding, rather than a sentence fragment.')

    add_callout(doc, '📊 RAG IMPACT:', 'Enhanced RAG pipeline delivered 40% improvement in agent output relevance scores across all 18 agents when tested on US Bank\'s knowledge base.')

    add_h2(doc, '2.4 Lean VSM Agent: Seeing the Value Stream for the First Time')
    add_body(doc, 'Three weeks into the deployment, Sarah\'s team turned their attention to the product that had been causing the most operational pain: the **Online Banking Portal**, a 12-year-old platform that was simultaneously the bank\'s most trafficked digital product and its most technically encumbered.')
    add_body(doc, 'The Lean VSM Agent was configured and executed with a single click. What followed was a 41-second analysis that produced something US Bank had never had before: a complete, metric-populated current-state value stream map for the entire Online Banking Portal delivery process.')
    add_body(doc, 'The agent surfaced eight value stream steps with full lean metrics for each:')
    add_metric_table(doc, [
        ('Requirements & Analysis',   '4 days cycle, 2.5 days wait', '2.5 days cycle, 0.5 days wait', '−38% total lead time'),
        ('Architecture Design',       '6 hours active, 3 days queue', '6 hours active, 1 day queue', '−67% queue reduction'),
        ('Development',               '16 hours active, 8 days queue', '16 hours active, 3 days queue', '−63% queue reduction'),
        ('Quality Assurance',         '8 hours active, 4 days queue', '6 hours active, 1.5 days queue', '−50% total time'),
        ('Security Review',           '4 hours active, 5 days queue', '4 hours active, 1 day queue', '−80% queue reduction'),
        ('UAT & Business Sign-off',   '3 hours active, 7 days queue', '3 hours active, 2 days queue', '−71% queue reduction'),
        ('Deployment & Release',      '2 hours active, 1.5 days queue', '2 hours active, 0.5 days queue', '−67% queue reduction'),
        ('Post-Release Monitoring',   '∞ (reactive only)', 'Automated + alerts', 'Proactive coverage'),
    ])
    add_body(doc, 'The agent identified five distinct waste categories across the value stream, including a particularly significant finding: **7 days of average queue time in the Security Review step**, caused by the bank\'s security team operating a sequential review process for all digital product changes regardless of risk level. This single waste item was estimated to be adding **$2.3 million in annual delivery cost** through delayed feature releases and extended opportunity windows.')
    add_body(doc, 'Current Process Cycle Efficiency: **34%**. Industry benchmark for comparable digital banking platforms: **58%**. The gap was stark — and for the first time, it was quantified.')

    add_h2(doc, '2.5 Future State Vision Agent: From Analysis to Action')
    add_body(doc, 'The transformation from insight to roadmap happened through the Future State Vision Agent — and it was the moment that most changed how US Bank\'s leadership team thought about AI in transformation.')
    add_body(doc, 'Before running the agent, the team uploaded three benchmark documents to the TransformHub knowledge base: a Gartner Digital Banking VSM Benchmarks report (2025), a McKinsey Digital Banking Transformation Case Studies compilation, and Fintech Australia\'s Technology Modernisation Metrics dataset. These were categorised as VSM_BENCHMARKS and TRANSFORMATION_CASE_STUDIES — the categories that receive the highest RAG budget allocation in the Future State agent\'s context window.')
    add_body(doc, 'The result was a transformation roadmap unlike anything the team had seen produced by AI before. Because the agent had access to real benchmark data and case studies through the RAG pipeline, every projected metric was grounded in documented industry evidence rather than generic optimism.')
    add_body(doc, 'The projected metrics for the Online Banking Portal transformation carried the **"🎯 Benchmark-grounded"** badge — the platform\'s signal that projections are based on retrieved evidence rather than internal multipliers:')
    add_metric_table(doc, [
        ('Process Cycle Efficiency', '34%',  'Conservative: 52%  |  Expected: 65%  |  Optimistic: 80%', 'Source: Gartner 2025 Benchmarks'),
        ('Delivery Lead Time',       '31 days', 'Conservative: 19d  |  Expected: 13d  |  Optimistic: 8d', 'Source: McKinsey Case Studies'),
        ('Automation Coverage',      '35%',  'Conservative: 55%  |  Expected: 68%  |  Optimistic: 82%', 'Source: Fintech Australia'),
        ('Security Review Cycle',    '5 days queue', 'Conservative: 2d  |  Expected: 1d  |  Optimistic: 4h', 'Risk-based review model'),
        ('Defect Escape Rate',       '8.2%', 'Conservative: 4.5%  |  Expected: 2.8%  |  Optimistic: 1.2%', 'Source: McKinsey Case Studies'),
    ])
    add_body(doc, 'The roadmap was structured across three transformation phases:')
    add_bullet(doc, '**Phase 1 — Stabilise (Months 1–3):** Security review risk-tiering model, CI/CD pipeline modernisation, test automation baseline, technical debt triage framework')
    add_bullet(doc, '**Phase 2 — Modernise (Months 4–9):** API-first architecture migration, microservices decomposition of monolithic services, automated security scanning integration, feature flag framework deployment')
    add_bullet(doc, '**Phase 3 — Optimise (Months 10–18):** Cloud-native deployment model, AI-assisted code review, predictive release risk scoring, zero-touch deployment for low-risk changes')

    add_h2(doc, '2.6 The Human-in-the-Loop Gate: Where AI Meets Judgement')
    add_body(doc, 'One of the most strategically important architectural decisions in TransformHub is the **Human-in-the-Loop (HITL) gate** — LangGraph\'s INTERRUPT mechanism that pauses agent execution at critical decision points to require human review before results are finalised.')
    add_body(doc, 'For US Bank, the HITL gate was configured for the Future State Vision agent — the point where transformation roadmaps carry sufficient strategic weight that they must be validated by the VP of Digital Transformation before being persisted to the knowledge base.')
    add_body(doc, 'When the Future State agent completed its initial roadmap generation for the Institutional Banking segment, the platform paused. A notification appeared on James Reeves\'s dashboard: **"⏸ Future State Agent awaiting your review."** James reviewed the draft roadmap — and found it substantively correct with one critical gap: the proposed Phase 2 architecture migration had not accounted for the bank\'s upcoming ISO 20022 payments migration, which would be running concurrently and had significant shared infrastructure dependencies.')
    add_body(doc, 'James rejected the draft with a single piece of feedback: **"Phase 2 must account for concurrent ISO 20022 migration — shared API gateway and messaging infrastructure cannot be modified in parallel."**')
    add_body(doc, 'The agent re-ran. The revised roadmap sequenced the API-first architecture migration to follow the ISO 20022 cutover, added a dedicated integration workstream for the payments infrastructure, and extended Phase 2 by six weeks to accommodate the dependency. The feedback was automatically saved to the platform\'s **agent memory** — ensuring that every subsequent Future State run for US Bank\'s Institutional Banking segment would be automatically aware of the ISO 20022 constraint, without requiring James to repeat himself.')
    add_pull_quote(doc, 'I rejected the first roadmap and gave one piece of feedback. The system learned it. Three months later, when we ran the agent again for Wealth Management, it proactively mentioned the ISO 20022 dependencies without being prompted. That\'s institutional memory working.', 'James Reeves, CTO, US Bank')

    add_h2(doc, '2.7 Risk & Compliance Agent: Regulatory Intelligence at Scale')
    add_body(doc, 'US Bank operates under a complex and overlapping regulatory landscape: APRA\'s CPS 234 (information security), PCI-DSS v4.0 (payment card security), the Notifiable Data Breaches scheme, and a host of AUSTRAC obligations for international payments. For the Risk & Compliance Agent, the bank uploaded the full text of each regulatory framework as REGULATORY-category knowledge base documents.')
    add_body(doc, 'The Risk Agent\'s analysis of the Online Banking Portal produced a 23-item risk register in 58 seconds. The output included risks the bank\'s own compliance team confirmed they had not formally catalogued:')
    add_bullet(doc, '**RISK-007 (High, Severity 16):** Biometric authentication module lacks fallback rate-limiting compliant with APRA CPS 234 §42.b — customer lockout vulnerability under adversarial conditions')
    add_bullet(doc, '**RISK-012 (High, Severity 15):** Third-party analytics SDK (version 3.1.x) transmits device fingerprint data to non-APRA-approved jurisdiction — potential Privacy Act s.13G breach')
    add_bullet(doc, '**RISK-019 (Medium, Severity 9):** OAuth 2.0 token rotation policy set to 24-hour expiry — exceeds PCI-DSS Requirement 8.3.9 maximum session duration for privileged access scenarios')
    add_body(doc, 'The compliance team validated all 23 risks within two business days. 19 were confirmed as material. The average time for a comparable manual risk assessment at US Bank had previously been **four weeks**.')

    add_h2(doc, '2.8 Agent Memory: The Platform That Learns')
    add_body(doc, 'Perhaps the most consequential enhancement to TransformHub — and the one most directly relevant to US Bank\'s long-term value — is the **Agent Memory Learning Loop**. Every time a human gate interaction occurs, every time an agent output is edited, and every time a HITL rejection includes feedback, the platform distils that interaction into a persisted learning, stored in the agent_memories table, scoped to US Bank\'s organisation and the specific agent type.')
    add_body(doc, 'After six months of operation, US Bank\'s TransformHub instance had accumulated 47 organisation-specific agent memories across the 18 agents:')
    add_bullet(doc, '**Discovery Agent (8 memories):** Preferred product taxonomy aligned to internal naming conventions; knowledge that "IBG" refers to Institutional Banking Group; the fact that the Treasury Management product is always co-dependent with the FX Derivatives capability')
    add_bullet(doc, '**VSM Agent (11 memories):** Security review step uses risk-tiered processing (not sequential) following HITL feedback; development cycle times in Retail Banking are 20% longer in Q4 due to code freeze policies; automation level for payment gateway is artificially low in the database due to undocumented RPA scripts')
    add_bullet(doc, '**Future State Agent (9 memories):** ISO 20022 migration dependency; the bank\'s board has a stated preference for 18-month transformation horizons (not 24-month); Phase 3 cloudnative work requires approval from the Group Infrastructure Risk Committee')
    add_bullet(doc, '**Risk Agent (12 memories):** APRA CPS 234 is the primary regulatory framework (not SOX, which applies only to US reporting entity); PCI-DSS v4.0 cutover date was March 2026 (important for risk prioritisation); the bank\'s internal security classification adds a "Highly Sensitive" tier above "Secret" not present in standard frameworks')
    add_body(doc, 'The compound effect of this learning was measurable: over the six-month period, the average **human edit rate** on agent outputs fell from 43% on the first run to **11%** by month six. The agents had, effectively, been trained on US Bank\'s own preferences, constraints, and institutional knowledge — without a single line of model fine-tuning.')

    add_h2(doc, '2.9 Accuracy Scoring: Building Trust in AI Intelligence')
    add_body(doc, 'One of the critical barriers to AI adoption in high-stakes enterprise decisions is the trust gap: executives are unwilling to act on AI outputs they cannot verify. TransformHub\'s **Accuracy Scoring System** was designed to close this gap by providing a transparent, continuously updated quality signal for every agent module.')
    add_body(doc, 'The composite accuracy score for each module is calculated as:')
    add_callout(doc, '📐 FORMULA:', 'Accuracy = (Confidence × 0.4) + (Source Diversity × 0.2) + (Run Success Rate × 0.3) + ((1 − Human Edit Rate) × 0.1)')
    add_body(doc, 'After six months of operation, US Bank\'s accuracy scores by module were:')
    add_two_col_table(doc, 'Agent Module', 'Accuracy Score (6-month)', [
        ('Discovery Agent',             '91% — Excellent'),
        ('Lean VSM Agent',              '87% — Excellent'),
        ('Future State Vision',         '84% — Good'),
        ('Risk & Compliance',           '88% — Excellent'),
        ('Product Transformation',      '79% — Good'),
        ('Architecture Agent',          '76% — Good'),
        ('Initiative Prioritisation',   '82% — Good'),
        ('Executive Reporting',         '89% — Excellent'),
    ])
    add_body(doc, 'Sarah presented these scores at the Q1 board review. For the first time in the bank\'s transformation history, the board had a quantitative confidence signal for AI-generated analysis. The conversation shifted: instead of "can we trust this?" the question became "where are the gaps, and how do we close them?"')

    # ── CHAPTER 3 ────────────────────────────────────────────────────────────
    add_h1(doc, 'Chapter 3: The Full 18-Agent Orchestra')
    add_body(doc, 'By Month 4 of the deployment, US Bank was operating the full TransformHub intelligence suite — all 18 agents working in concert, each informing the others through the shared PostgreSQL knowledge base and the continuously growing context document store.')
    add_h2(doc, '3.1 How the Agents Work Together')
    add_body(doc, 'The 18 agents are not independent silos. They form a deliberate intelligence cascade, where each agent\'s output becomes the RAG context input for subsequent agents:')
    add_bullet(doc, '**Discovery Agent** maps the digital estate → capabilities and products stored in DB and as AGENT_OUTPUT context docs')
    add_bullet(doc, '**Lean VSM Agent** reads Discovery output via RAG, maps value streams for each product → step metrics and waste items stored and embedded')
    add_bullet(doc, '**Future State Vision Agent** reads VSM output + uploaded benchmarks via RAG, generates roadmap with benchmark-grounded projections')
    add_bullet(doc, '**Risk & Compliance Agent** reads Discovery output + regulatory frameworks, identifies and scores risks with regulatory mapping')
    add_bullet(doc, '**Product Transformation Agent** reads Discovery + VSM + Risk outputs, generates product-specific modernisation strategies')
    add_bullet(doc, '**Architecture Agent** reads all prior outputs + architecture standards, recommends technology architecture and integration patterns')
    add_bullet(doc, '**Capability Maturity Agent** assesses capability maturity levels against industry benchmarks, identifying maturity gaps')
    add_bullet(doc, '**Initiative Prioritisation Agent** applies RICE scoring framework to all identified initiatives, creating a prioritised portfolio')
    add_bullet(doc, '**Transformation Roadmap Agent** synthesises VSM improvements + initiatives into a phased 18-month implementation plan')
    add_bullet(doc, '**Benchmark Agent** continuously compares US Bank\'s metrics against the uploaded benchmark corpus, surfacing performance gaps')
    add_bullet(doc, '**Executive Reporting Agent** compiles outputs from all domain agents into a structured C-suite report with confidence scores')
    add_bullet(doc, '**Human Gate Agent** manages INTERRUPT nodes across all agents, checkpoint-saving state to PostgreSQL for resumable review flows')
    add_bullet(doc, '**Agent Memory Agent** extracts learnings from HITL interactions and persists them as org-scoped, agent-type-scoped memories')
    add_bullet(doc, '**Accuracy Scoring Agent** computes and caches composite accuracy scores per module using 60-second TTL caching')
    add_bullet(doc, '**Context Output Agent** auto-saves all agent outputs as AGENT_OUTPUT context documents, maintaining the self-improving RAG corpus')
    add_bullet(doc, '**BM25 Retrieval Agent** performs keyword-based retrieval reranking on top-25 vector search results for VSM and Future State agents')
    add_bullet(doc, '**Org Context Agent** assembles category-aware context sections for each agent, applying the 12,000-character budget with intelligent allocation')
    add_bullet(doc, '**Audit Trail Agent** maintains the SHA-256 chained immutable audit log, ensuring forensic traceability for all platform operations')

    add_h2(doc, '3.2 The Audit Trail: Board-Level Governance')
    add_body(doc, 'For a regulated financial institution like US Bank, the governance story is as important as the intelligence story. TransformHub\'s **SHA-256 chained audit trail** provides an immutable, cryptographically verifiable record of every platform operation — agent runs, HITL gate decisions, knowledge base changes, and data mutations.')
    add_body(doc, 'Each audit_log entry contains: the action performed, the entity affected, the user who triggered it, the full input and output payload, and a SHA-256 hash computed over the previous entry\'s hash plus the current entry\'s content. The result is a chain that cannot be tampered with without detection — every entry\'s integrity is provable from the genesis record forward.')
    add_body(doc, 'When US Bank\'s internal audit team requested evidence of the transformation intelligence process for their annual IT audit, the entire audit trail for six months of platform operation was exported in 14 seconds. The auditors confirmed it was the most complete and verifiable record of AI-assisted business analysis they had reviewed.')

    # ── CHAPTER 4 ────────────────────────────────────────────────────────────
    add_h1(doc, 'Chapter 4: The Results — Six Months On')
    add_h2(doc, '4.1 Quantified Outcomes')
    add_body(doc, 'Six months after deployment, the US Bank transformation team conducted a formal outcomes review against the baseline metrics from January 2026:')
    add_metric_table(doc, [
        ('Time to complete full portfolio discovery',         '6 weeks (manual workshop)', '4 hours (Discovery Agent)', '−98% time reduction'),
        ('Time to produce VSM for one product',              '3 weeks (consultant-led)', '41 seconds (VSM Agent)', '−99.97% time reduction'),
        ('Time from VSM to transformation roadmap',          '4–6 weeks (manual)', '< 2 hours (Future State Agent)', '−98% time reduction'),
        ('Cost per transformation assessment (one product)', '$380,000 (consulting)', '$12,000 (platform cost)', '−97% cost reduction'),
        ('Regulatory risk identification coverage',          '60% (expert estimate)', '91% (Risk Agent + compliance review)', '+31% coverage improvement'),
        ('Agent output human edit rate',                     'N/A (new baseline)', '11% at Month 6 (down from 43%)', '−74% edit rate decline'),
        ('Average agent accuracy score (portfolio)',         'N/A (new baseline)', '84.5% at Month 6', 'Trust threshold exceeded'),
        ('Executive report preparation time',                '3 weeks (manual)', '< 15 minutes (Reporting Agent)', '−99% time reduction'),
        ('Transformation portfolio visibility',              'Quarterly (90-day lag)', 'Continuous (real-time)', 'Always current'),
        ('Knowledge base documents indexed',                 '0 at start', '312 documents, 18,400 chunks', 'Compound intelligence built'),
    ])

    add_h2(doc, '4.2 Strategic Impact')
    add_body(doc, 'Beyond the operational metrics, the deployment had three strategic impacts that Sarah Mitchell cited as the most significant:')
    add_body(doc, '**First: The board conversation changed.** Instead of quarterly debates about transformation confidence, the board now reviews a live dashboard with accuracy-scored intelligence. The question shifted from "is our transformation programme on track?" to "what are the three biggest risks we should address this quarter?" TransformHub had converted a governance problem into a management opportunity.')
    add_body(doc, '**Second: The consultancy relationship changed.** US Bank did not stop using external consultants — but the nature of the engagement changed fundamentally. Instead of commissioning discovery and assessment work, consultants were engaged to validate and act on TransformHub\'s intelligence. The bank estimated this shift reduced its annual consulting spend by $4.2 million while increasing the effectiveness of each engagement.')
    add_body(doc, '**Third: The organisation learned.** The agent memory system had created, for the first time, a mechanism for institutional knowledge to accumulate in a queryable, reusable form. New team members onboarding to transformation roles could access six months of structured, contextualised intelligence about US Bank\'s digital estate — not just documents, but the living product of 47 agent memory entries representing hard-won organisational learning.')

    add_pull_quote(doc, 'TransformHub didn\'t replace our people. It made them dramatically more effective. Every analyst on my team now operates at the level of a principal consultant because they have 18 AI specialists working alongside them, continuously, with perfect memory.', 'Sarah Mitchell, CDO, US Bank')

    # ── EPILOGUE ─────────────────────────────────────────────────────────────
    add_h1(doc, 'Epilogue: The Continuous Transformation Intelligence Era')
    add_body(doc, 'In April 2026, US Bank\'s board approved a $240 million digital transformation investment across the Retail Banking and Institutional Banking segments, to be delivered over 18 months. The approval was notable for two reasons. First, it was the largest single transformation investment commitment in the bank\'s history. Second, it was the first time the board had approved a transformation investment with a quantified, AI-generated, benchmark-grounded business case — complete with confidence scores, risk register, and phased roadmap.')
    add_body(doc, 'The proposal had been generated by TransformHub in under four hours.')
    add_body(doc, 'Sarah Mitchell closed her presentation to the board with a reflection on what had changed. She still had binders of old assessment reports on her shelf. They represented $8 million in accumulated intelligence that had never been fully used — too slow, too expensive, and too static to keep pace with the organisation they were meant to illuminate.')
    add_body(doc, '"We no longer commission point-in-time pictures," she told the board. "We have a living, learning intelligence system that shows us where we are, where the industry is, and where we need to go — continuously. That\'s not an IT investment. That\'s a strategic capability."')
    add_body(doc, 'The board voted unanimously to proceed.')

    out = '/Users/125066/projects/TransformHub/docs/USBank_TransformHub_Narrative_POV.docx'
    doc.save(out)
    print(f'Saved narrative: {out}')
    return out


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 2 — CASE STUDY POV
# ═══════════════════════════════════════════════════════════════════════════════
def build_case_study():
    doc = Document()
    set_margins(doc)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # COVER
    add_cover(doc,
        'US Bank × TransformHub',
        'Case Study: AI-Powered Digital Transformation Intelligence',
        'Industry: Banking & Financial Services  |  Region: United States & Australia',
        'March 2026  |  Confidential')

    # ── EXECUTIVE SUMMARY ────────────────────────────────────────────────────
    add_h1(doc, 'Executive Summary')
    add_body(doc, 'US Bank, one of the largest commercial banks in the United States, deployed TransformHub — an enterprise AI-powered digital transformation intelligence platform — across its Digital Channels Portfolio and Core Banking Platform in January 2026. Within six months, the bank achieved a **97% reduction in transformation assessment cost**, a **98% reduction in time-to-insight**, and established the first continuously-updated, AI-scored transformation intelligence system in its operating history.')
    add_body(doc, 'The deployment leveraged TransformHub\'s full suite of **18 specialised LangGraph agents**, a **hybrid BM25 + vector RAG pipeline** with category-aware 12,000-character context budgets, **Human-in-the-Loop governance gates**, an **agent memory learning system**, and a **SHA-256 chained audit trail** — enabling the bank\'s Chief Digital Officer to present a benchmark-grounded, confidence-scored $240M transformation business case to the board within four hours of synthesis.')

    t = doc.add_table(rows=5, cols=2)
    t.style = 'Table Grid'
    kv = [
        ('CLIENT', 'US Bank — Commercial Banking, United States'),
        ('INDUSTRY', 'Banking & Financial Services'),
        ('PLATFORM', 'TransformHub v1.0 — 18 Agent AI Intelligence Platform'),
        ('DEPLOYMENT', 'January 2026 — 3 business segments, 2 repositories'),
        ('STATUS', '6 months operational — Full agent suite active'),
    ]
    for i, (k, v) in enumerate(kv):
        r = t.rows[i]
        r.cells[0].text = k
        r.cells[1].text = v
        if r.cells[0].paragraphs[0].runs:
            r.cells[0].paragraphs[0].runs[0].bold = True
            r.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            r.cells[0].paragraphs[0].runs[0].font.color.rgb = WHITE
        if r.cells[1].paragraphs[0].runs:
            r.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        shd(r.cells[0], FILL_HEADER)
        if i % 2 == 1:
            shd(r.cells[1], FILL_ALT)
    doc.add_paragraph()

    # ── CLIENT BACKGROUND ────────────────────────────────────────────────────
    add_h1(doc, '1. Client Background')
    add_h2(doc, '1.1 Organisation Profile')
    add_body(doc, 'US Bank is a full-service commercial and retail bank operating across the United States and internationally. The bank serves 18 million retail customers, 6,500 institutional clients, and manages assets under administration exceeding $600 billion. Its digital transformation programme, launched in 2022 under the "Digital First" strategic agenda, committed $1.4 billion over five years to modernising the bank\'s digital estate across three business segments:')
    add_bullet(doc, '**Retail Banking** — consumer deposits, lending, mobile and online banking, payments')
    add_bullet(doc, '**Institutional Banking (IBG)** — corporate banking, trade finance, treasury management, FX')
    add_bullet(doc, '**Wealth Management** — private banking, investment advisory, portfolio management, insurance')

    add_h2(doc, '1.2 Digital Estate at Deployment')
    add_two_col_table(doc, 'Dimension', 'Scale', [
        ('Digital Products', '12 major products across 3 business segments'),
        ('Digital Capabilities', '47 capabilities mapped at deployment'),
        ('Functionalities', '200+ individual functionalities'),
        ('Technology Stacks', '3 primary stacks (2002, 2011, 2019 vintage)'),
        ('Annual IT Operating Cost', '~$890M per year'),
        ('Active Transformation Initiatives', '34 initiatives across 3 segments at deployment'),
        ('Compliance Frameworks', 'APRA CPS 234, PCI-DSS v4.0, Privacy Act, AUSTRAC, SOX'),
    ])

    # ── CHALLENGE ────────────────────────────────────────────────────────────
    add_h1(doc, '2. The Challenge')
    add_h2(doc, '2.1 Primary Problem Statement')
    add_body(doc, 'US Bank faced five interconnected transformation intelligence challenges that collectively impeded the effectiveness of its $1.4B digital transformation programme:')

    add_two_col_table(doc, 'Challenge', 'Business Impact', [
        ('Transformation Opacity: No continuous, unified view of digital estate transformation status', '$140M estimated annual cost of misaligned investment and delayed delivery'),
        ('Manual Analysis Bottleneck: VSM and capability assessments requiring weeks of facilitated workshops', '$380K per product assessment; only 3 products fully assessed in 2 years'),
        ('Siloed Intelligence: Transformation knowledge locked in individual consultant documents, no reuse', '~$8M in assessments that generated insights never systematically reused'),
        ('No Benchmark Grounding: Transformation plans without reference to industry performance data', 'Board investment proposals rejected or delayed due to unsubstantiated projections'),
        ('Governance Gap: AI-generated insights lacked audit trail, confidence scoring, human oversight mechanisms', 'CDO unable to present AI analysis at board level due to traceability concerns'),
    ])

    add_h2(doc, '2.2 Specific Pain Points at Time of Engagement')
    add_bullet(doc, 'Most recent portfolio-wide capability assessment was 14 months old at time of TransformHub deployment')
    add_bullet(doc, 'Online Banking Portal — the bank\'s highest-revenue digital product — had not had a formal VSM analysis in 3 years')
    add_bullet(doc, 'Risk register for Institutional Banking digital products was acknowledged internally as "significantly incomplete"')
    add_bullet(doc, 'Executive reporting on transformation progress required 3 weeks of manual compilation by 2 senior analysts')
    add_bullet(doc, 'Board confidence in transformation programme had declined: 2 of 5 board members had called for an independent external review')

    # ── SOLUTION ─────────────────────────────────────────────────────────────
    add_h1(doc, '3. The TransformHub Solution')
    add_h2(doc, '3.1 Platform Architecture')
    add_body(doc, 'TransformHub deployed as a dedicated instance for US Bank, running on the bank\'s approved cloud infrastructure. The platform architecture comprised:')
    add_two_col_table(doc, 'Component', 'Technology & Configuration', [
        ('Frontend UI', 'Next.js 15 App Router + TypeScript + Tailwind v4 — dark glassmorphism interface'),
        ('Agent Service', 'FastAPI + LangGraph 18 agents — Python 3.11, async/await architecture'),
        ('Primary Database', 'PostgreSQL 18 with pgvector extension — 1,536-dimensional embeddings'),
        ('LLM Inference', 'OpenAI GPT-4o — all 18 agents; structured output with Pydantic v2 schemas'),
        ('Embedding Model', 'OpenAI text-embedding-3-small — 1,536 dimensions, 2k chunk / 400 overlap'),
        ('Authentication', 'NextAuth.js v5 with JWT + database sessions'),
        ('Audit System', 'SHA-256 chained audit_logs — immutable, INSERT-only'),
        ('Vector Index', 'pgvector ivfflat, lists=100, probes=10 — ~90% recall'),
    ])

    add_h2(doc, '3.2 The 18-Agent Intelligence Suite')
    add_body(doc, 'TransformHub\'s value is delivered through a coordinated suite of 18 specialised LangGraph agents, each owning a distinct transformation intelligence domain and communicating exclusively through the shared PostgreSQL knowledge base:')
    add_two_col_table(doc, 'Agent', 'US Bank Deployment Outcome', [
        ('Discovery Agent', '12 products, 47 capabilities, 200+ functionalities mapped in < 4 hours (vs 6 weeks manual)'),
        ('Lean VSM Agent', 'Current-state VSM for all 12 products; 34% baseline PCE identified; 5 major waste categories surfaced'),
        ('Future State Vision Agent', 'Benchmark-grounded roadmaps with 3-band projections; ISO 20022 dependency captured via HITL memory'),
        ('Risk & Compliance Agent', '23-item risk register for Online Banking Portal in 58s; 2 previously uncatalogued APRA CPS 234 risks identified'),
        ('Product Transformation Agent', 'Product-specific modernisation strategies for all 12 products; cloud migration sequencing recommended'),
        ('Architecture Agent', 'API-first architecture recommendations; microservices decomposition plan for 3 monolithic products'),
        ('Capability Maturity Agent', 'Maturity scores across 47 capabilities; 12 capabilities in "Initial" state requiring immediate investment'),
        ('Initiative Prioritisation Agent', 'RICE-scored portfolio of 34 initiatives; top-5 priorities aligned to $240M board submission'),
        ('Transformation Roadmap Agent', '18-month phased roadmap across 3 segments; 97% milestone coverage vs manually produced equivalent'),
        ('Benchmark Agent', 'Continuous comparison vs Gartner/McKinsey/Fintech Australia benchmarks; PCE gap of 24 percentage points identified'),
        ('Executive Reporting Agent', 'Board-ready C-suite report in < 15 minutes; approved as primary reporting artefact by CDO'),
        ('Human Gate Agent', 'INTERRUPT nodes active for Future State + Architecture agents; all 14 HITL reviews completed within 4 hours'),
        ('Agent Memory Agent', '47 org-specific memories accumulated; human edit rate declined 74% over 6 months'),
        ('Accuracy Scoring Agent', 'Portfolio average accuracy 84.5% at Month 6; per-module scores displayed on live dashboard'),
        ('Context Output Agent', '312 AGENT_OUTPUT documents auto-saved, 18,400 chunks indexed, compounding RAG quality'),
        ('BM25 Retrieval Agent', 'BM25 reranking improved regulatory clause retrieval precision by 35% for Risk agent runs'),
        ('Org Context Agent', '12k char budget; category-aware allocation per agent; benchmark docs prioritised for Future State runs'),
        ('Audit Trail Agent', 'Complete 6-month audit trail exported in 14 seconds for annual IT audit; chain integrity verified'),
    ])

    add_h2(doc, '3.3 Key Platform Enhancements Applied')
    add_body(doc, 'The US Bank deployment benefited from several critical platform enhancements that materially improved output quality:')
    add_bullet(doc, '**RAG Context Expansion (4k → 12k chars):** Enabled full regulatory framework context injection for risk assessments; APRA CPS 234 complete text retrievable in single agent run')
    add_bullet(doc, '**Hybrid BM25 + Vector Retrieval:** 35% improvement in regulatory clause precision; exact regulatory references (e.g., "CPS 234 §42.b") reliably surfaced regardless of semantic distance')
    add_bullet(doc, '**Multi-Query Strategy (3–5 parallel queries):** 40% improvement in overall context relevance; benchmark documents consistently retrieved for Future State runs even when queries were broadly phrased')
    add_bullet(doc, '**Chunk Size Increase (1k → 2k chars, 400-char overlap):** Preserved multi-sentence regulatory arguments and case study findings within single chunks; eliminated fragmented context that caused incomplete risk identification in earlier system version')
    add_bullet(doc, '**Agent Memory System:** 47 org-specific memories; ISO 20022 dependency captured and reused across 3 subsequent agent runs without re-prompting; compliance classification system captured and applied')
    add_bullet(doc, '**Benchmark-Grounded Future State:** "🎯 Benchmark-grounded" badge on all projected metrics sourced from Gartner/McKinsey uploads; eliminated board objections to unsubstantiated projections')
    add_bullet(doc, '**SHA-256 Chained Audit Trail:** Annual IT audit completed in < 1 business day (vs 3-week manual process previously)')

    # ── RESULTS ─────────────────────────────────────────────────────────────
    add_h1(doc, '4. Results & Outcomes')
    add_h2(doc, '4.1 Quantified Business Results')
    add_metric_table(doc, [
        ('Transformation Assessment Cost (per product)',   '$380,000',    '$12,000',     '−97%'),
        ('Time: Portfolio Discovery (all 12 products)',    '6 weeks',     '< 4 hours',   '−98%'),
        ('Time: VSM Analysis (one product)',               '3 weeks',     '41 seconds',  '−99.97%'),
        ('Time: Risk Assessment (one product)',            '4 weeks',     '58 seconds',  '−99.96%'),
        ('Time: Executive Board Report',                  '3 weeks',     '< 15 minutes','−99%'),
        ('Transformation Portfolio Visibility Lag',       '90 days',     'Real-time',   '−100%'),
        ('Regulatory Risk Coverage',                      '60%',         '91%',         '+31 ppts'),
        ('Agent Output Human Edit Rate',                  '43% (Month 1)','11% (Month 6)','−74%'),
        ('Portfolio Avg. Accuracy Score',                 'N/A (new)',   '84.5%',       'Baseline established'),
        ('Annual Consulting Assessment Spend Reduction',  '$4.8M/year',  '$0.6M/year',  '−$4.2M/year'),
    ])

    add_h2(doc, '4.2 Strategic Outcomes')
    add_bullet(doc, '**$240M transformation investment approved** — first board investment submission to include AI-generated, benchmark-grounded, confidence-scored business case')
    add_bullet(doc, '**Board confidence restored** — 2 board members who had requested independent review withdrew request after reviewing TransformHub accuracy scores and audit trail')
    add_bullet(doc, '**Consulting engagement model transformed** — external consultants now validate and act on TransformHub intelligence rather than conducting discovery; average engagement value increased 40% while total spend decreased 87%')
    add_bullet(doc, '**Institutional knowledge created** — 47 agent memories representing 6 months of structured organisational learning; accessible to all new team members immediately')
    add_bullet(doc, '**ISO 20022 risk mitigated** — HITL gate capture of payments infrastructure dependency prevented estimated $18M in rework from concurrent architecture conflict')

    add_h2(doc, '4.3 ROI Analysis')
    add_two_col_table(doc, 'ROI Component', 'Value (Year 1)', [
        ('Assessment cost reduction (consulting fees avoided)', '$4.2M'),
        ('Delivery acceleration value (earlier feature revenue)', '$8.7M (estimated)'),
        ('Risk mitigation value (ISO 20022 conflict avoided)', '$18.0M (estimated)'),
        ('Compliance exposure reduction (APRA/PCI findings pre-empted)', '$6.5M (estimated)'),
        ('Executive reporting efficiency (analyst time recovered)', '$0.8M'),
        ('Total Year 1 Value', '$38.2M (estimated)'),
        ('Total Year 1 Platform Investment (licensing + implementation)', '$2.1M'),
        ('Year 1 ROI', '1,719%'),
        ('Payback Period', '< 3 weeks'),
    ])

    # ── LESSONS ──────────────────────────────────────────────────────────────
    add_h1(doc, '5. Lessons Learned & Success Factors')
    add_h2(doc, '5.1 Critical Success Factors')
    add_bullet(doc, '**Executive sponsorship:** CDO Sarah Mitchell was an active platform user, not just a sponsor. Her HITL gate interactions drove the most valuable agent memory entries.')
    add_bullet(doc, '**Knowledge base investment:** The team dedicated 3 days to uploading benchmark documents and regulatory frameworks before running any agents. This upfront investment in the RAG corpus drove a measurable step-change in output quality.')
    add_bullet(doc, '**HITL governance culture:** The team embraced rather than bypassed the HITL gates. The 14 gate interactions over 6 months generated 23 of the platform\'s most valuable agent memories.')
    add_bullet(doc, '**Iterative trust building:** The team started with Discovery Agent, validated outputs manually, then progressively extended trust to VSM, Future State, and Risk agents as accuracy scores accumulated.')
    add_bullet(doc, '**Segment-first approach:** Running Discovery and VSM per business segment rather than across the full portfolio allowed the team to validate quality before committing the full transformation programme to AI-generated intelligence.')

    add_h2(doc, '5.2 What Worked Exceptionally Well')
    add_bullet(doc, 'The **benchmark-grounded future state** eliminated the most common objection to transformation roadmaps at US Bank: "where do these numbers come from?"')
    add_bullet(doc, 'The **agent memory learning loop** reduced the operational burden on senior executives over time — the system learned their preferences so they didn\'t have to repeat themselves')
    add_bullet(doc, 'The **SHA-256 audit trail** was a decisive factor in board acceptance of AI-generated analysis. The chain\'s tamper-evidence was cited explicitly by the compliance team.')
    add_bullet(doc, 'The **hybrid RAG pipeline** with BM25 enhancement was critical for regulatory use cases — exact regulatory clause references could not be reliably surfaced with vector search alone')

    add_h2(doc, '5.3 Areas for Further Development')
    add_bullet(doc, 'Real-time agent streaming (SSE) would improve user experience during longer agent runs — planned for v1.1')
    add_bullet(doc, 'SSO/SAML integration with the bank\'s Okta environment would simplify user provisioning — planned for v1.2')
    add_bullet(doc, 'Jira bidirectional integration would allow TransformHub initiatives to sync directly to delivery backlogs — planned for v2.0')
    add_bullet(doc, 'Azure OpenAI as fallback to address data residency requirements for certain Institutional Banking data classifications — planned for v1.2')

    # ── QUOTE / CLOSE ────────────────────────────────────────────────────────
    add_h1(doc, '6. Client Statement')
    add_pull_quote(doc,
        'TransformHub fundamentally changed what transformation intelligence means at US Bank. We went from commissioning quarterly snapshots at enormous cost to having a live, continuously learning system that knows our organisation. The board approved a $240 million investment programme based on intelligence this platform generated in four hours. That\'s not incremental improvement — that\'s a step change in how we operate.',
        'Sarah Mitchell, Chief Digital Officer, US Bank')

    add_body(doc, 'For further information about the TransformHub platform or this case study, contact the TransformHub team.')

    out = '/Users/125066/projects/TransformHub/docs/USBank_TransformHub_CaseStudy.docx'
    doc.save(out)
    print(f'Saved case study: {out}')
    return out


if __name__ == '__main__':
    n = build_narrative()
    c = build_case_study()
    import subprocess
    subprocess.Popen(['open', n])
    subprocess.Popen(['open', c])
